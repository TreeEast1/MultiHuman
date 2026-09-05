#!/usr/bin/env python3
"""把方法流程图、数据流图写入 S 实验报告 §1.2，并按报告模板改页边距。"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

sys.path.insert(0, "/tmp")
from report_fmt import SONG, _force_embedded_inline, add_run, set_run_font

SRC = Path(
    "/Users/licochen/Library/Containers/com.tencent.xinWeChat/Data/Documents/"
    "xwechat_files/wxid_z6fnqx80s5z332_ef48/temp/drag/"
    "基于多模态生理与行为信号的操纵员绩效 S 评估实验报告.docx"
)
FIG = Path(__file__).resolve().parent / "figures_flow"
DESK = Path("/Users/licochen/Desktop")
OUT_NAME = "基于多模态生理与行为信号的操纵员绩效 S 评估实验报告.docx"


def insert_after(anchor: Paragraph) -> Paragraph:
    new_el = OxmlElement("w:p")
    anchor._p.addnext(new_el)
    return Paragraph(new_el, anchor._parent)


def style_body(p: Paragraph, text: str) -> None:
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, text, SONG, 12)


def style_caption(p: Paragraph, text: str) -> None:
    p.style = "图注"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    add_run(p, text, SONG, 10.5)


def style_figure(p: Paragraph, image: Path, width_cm: float) -> None:
    p.style = "图"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:before"), "120")
    sp.set(qn("w:after"), "40")
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "auto")
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:firstLineChars"), "0")
    ind.set(qn("w:left"), "0")
    run = p.add_run()
    run.add_picture(str(image), width=Cm(width_cm))
    _force_embedded_inline(run)


def set_text(p: Paragraph, text: str, size=12) -> None:
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], SONG, size)
        for r in p.runs[1:]:
            r.text = ""
    else:
        add_run(p, text, SONG, size)


def main() -> None:
    work = Path("/tmp") / OUT_NAME
    shutil.copy2(SRC, work)
    doc = Document(str(work))

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(1.75)

    # 后文图号后移：先改 5→6，避免覆盖
    replacements = [
        ("如图 5 所示", "如图 6 所示"),
        ("图 5  模态消融（合成 S 的决定系数）", "图 6  模态消融（合成 S 的决定系数）"),
        ("如图 4 所示", "如图 5 所示"),
        ("图 4  公式法下各算法的合成 S", "图 5  公式法下各算法的合成 S"),
        ("图 3 给出各折柱状对比", "图 4 给出各折柱状对比"),
        ("图 3  五折中合成 S 的决定系数", "图 4  五折中合成 S 的决定系数"),
        ("图 2 给出 84 条真值与预测的散点", "图 3 给出 84 条真值与预测的散点"),
        ("图 2  真值 S 与预测 S（84 条，正式口径）", "图 3  真值 S 与预测 S（84 条，正式口径）"),
    ]
    for p in doc.paragraphs:
        t = p.text or ""
        for old, new in replacements:
            if old in t:
                set_text(p, t.replace(old, new), size=10.5 if p.style.name == "图注" else 12)
                t = p.text

    body_p = fig_p = cap_p = after_p = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("技术路线为：四模态连续信号"):
            body_p = p
        elif t.startswith("图 1  预测绩效 S 的技术路线") or t.startswith("图 1  全模态绩效"):
            cap_p = p
        elif t.startswith("信号、切窗、降维、回归与公式合成"):
            after_p = p
        elif body_p is not None and fig_p is None and "w:drawing" in p._p.xml:
            fig_p = p

    if body_p is None or fig_p is None or cap_p is None or after_p is None:
        raise RuntimeError(f"anchors missing: {body_p, fig_p, cap_p, after_p}")

    set_text(
        body_p,
        "技术路线为：四模态连续信号按 30 秒窗、5 秒步切分，提取 66 个窗口指标；"
        "每个指标在任务内再汇总为均值、标准差、中位数与斜率，得到 84×264 的任务级表；"
        "每一折仅在训练被试上按模态定额做互信息筛选，得到 27 维；用浅树 XGBoost 估计负荷；"
        "再与真实步骤分合成 S。划分方式为按被试的五折 GroupKFold，同一人不同时出现在训练与测试。",
    )

    fig_p._element.getparent().remove(fig_p._element)
    cap_p._element.getparent().remove(cap_p._element)

    cursor = body_p
    p = insert_after(cursor)
    style_figure(p, FIG / "fig1_s_method_flow.png", 14.2)
    cursor = p
    p = insert_after(cursor)
    style_caption(p, "图 1  全模态绩效 S 预测方法流程图")
    cursor = p
    p = insert_after(cursor)
    style_body(
        p,
        "图 1 给出从原始数据到预测绩效 Ŝ 的计算顺序：四模态时序、NASA-TLX 问卷与步骤完成表经切窗与任务级汇总得到 84×264，"
        "按被试五折在训练折内定额至 27 维，浅树 XGBoost 估计负荷；五折拼合后，与真实步骤分按 0.70／0.30 合成 Ŝ，再与真值 S 对照。",
    )
    cursor = p
    p = insert_after(cursor)
    style_figure(p, FIG / "fig2_s_data_flow.png", 15.2)
    cursor = p
    p = insert_after(cursor)
    style_caption(p, "图 2  全模态绩效 S 预测数据流图")
    cursor = p
    p = insert_after(cursor)
    style_body(
        p,
        "图 2 给出同一路径上的数据对象：左侧由四模态信号得到窗级 n_win×66、任务级 84×264 与定额后 84×27，并输出折外负荷预测 L̂；"
        "右侧由问卷与步骤表得到 L、S_step 与真值 S；L̂ 与真实 S_step 合成 Ŝ，再与真值对照得到 R² 与 MAE。",
    )

    out = DESK / OUT_NAME
    doc.save(out)
    print("wrote", out, out.stat().st_size)
    try:
        shutil.copy2(out, SRC)
        print("also wrote source", SRC)
    except OSError as e:
        print("skip source copy:", e)


if __name__ == "__main__":
    main()
