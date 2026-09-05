#!/usr/bin/env python3
"""把 α 曲面图与 α=0.70 下的 S 结果整理成一份可直接外发的 Word。

依赖 plot_s_alpha_surface.py 先产出 figures/*_docx.png。

运行：
    uv run --with pandas --with numpy --with python-docx --with scikit-learn \
        python make_docx_s_alpha.py
"""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sklearn.metrics import mean_absolute_error, r2_score

HERE = Path(__file__).resolve().parent
FIG_DIR = HERE / "figures"
S_TABLE = HERE / "output" / "s_score_84samples.csv"
PRED_DIR = HERE / "reports_s_fullmodal"
OUT_DOCX = HERE / "给翁子伟_合成绩效S的α曲面与结果.docx"

ALPHA_MAIN = 0.70
ALPHA_PANELS = (0.30, 0.50, 0.70)
TASK_ORDER = ["1", "2", "3", "4", "5", "5_6"]
BODY_CN = "宋体"
BODY_EN = "Times New Roman"
HEAD_CN = "黑体"
FIG_WIDTH = Cm(16.0)


def set_run_font(run, cn: str = BODY_CN, en: str = BODY_EN, size: float = 10.5,
                 bold: bool = False) -> None:
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)


def add_para(doc, text: str, size: float = 10.5, cn: str = BODY_CN, bold: bool = False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent: bool = True, space_after: float = 6):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.25
    if indent:
        pf.first_line_indent = Pt(21)
    set_run_font(p.add_run(text), cn=cn, size=size, bold=bold)
    return p


def add_heading(doc, text: str, size: float = 13) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run(text), cn=HEAD_CN, size=size, bold=True)


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run(text), cn=BODY_CN, size=9.5)


def add_figure(doc, path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=FIG_WIDTH)


def add_formula(doc, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_run_font(r, cn=BODY_CN, en="Cambria Math", size=11)


def add_table(doc, header: list[str], rows: list[list[str]], widths: list[float]) -> None:
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, name in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.width = Cm(widths[j])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(name), cn=HEAD_CN, size=10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].width = Cm(widths[j])
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(val), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build() -> Path:
    df = pd.read_csv(S_TABLE)
    df["task"] = df["task"].astype(str)
    step = df["weighted_step_score"].to_numpy(dtype=float)
    nasa_rev = df["nasa_reverse"].to_numpy(dtype=float)
    s07 = ALPHA_MAIN * step + (1 - ALPHA_MAIN) * nasa_rev
    df["S07"] = s07

    y_true = np.load(PRED_DIR / "y_s07.npy")
    y_hat = np.load(PRED_DIR / "yhat_s07_xgb.npy")
    pred_r2 = r2_score(y_true, y_hat)
    pred_mae = mean_absolute_error(y_true, y_hat)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_EN
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run("合成绩效 S 的 α 曲面与本实验结果"), cn=HEAD_CN, size=16, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("26 名被试 / 84 次被试–任务　·　正式口径 α = 0.70")
    set_run_font(r, size=10.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_heading(doc, "一、公式与 α 的含义")
    add_para(
        doc,
        "综合绩效 S 由客观的步骤分与主观的 NASA-TLX 反向分线性合成，α 为客观步骤分的权重。"
        "本实验的正式口径取 α = 0.70。",
    )
    add_formula(doc, [
        "步骤分 = 0.75 × 关键子任务完成率 + 0.25 × 非关键子任务完成率",
        "NASA 反向分 = 1 − NASA-TLX 加权总分 / 10",
        "S = α × 步骤分 + (1 − α) × NASA 反向分，　α = 0.70",
    ])
    add_para(
        doc,
        "需要说明一处记号差异：您示意图中的 α 是主观值 x₁ 的权重，而我们报告中的 α = 0.70 指的是"
        "客观步骤分的权重。下图按我们报告的定义绘制，因此图题写作「α 越大，客观步骤分的贡献越大」，"
        "两个自变量轴分别标注为「客观值 · 步骤分」与「主观值 · NASA 反向分」。视角、配色与三联布局"
        "与您给的示意图保持一致。",
    )

    add_heading(doc, "二、不同 α 下的 S 曲面与真实样本位置")
    add_para(
        doc,
        "由于 S 是两个分量的线性合成，曲面在任意 α 下都是一张平面：α 越大平面沿步骤分方向越陡，"
        "沿 NASA 反向分方向越平。图中红点为本实验 84 条被试–任务样本的实际取值，按定义严格落在曲面上，"
        "因此可以直接读出真实数据在整张平面上所占的区域。",
    )
    add_figure(doc, FIG_DIR / "fig1_s_alpha_surface_docx.png")
    add_caption(
        doc,
        "图 1　不同 α 下的合成绩效 S 曲面（α = 0.30 / 0.50 / 0.70），红点为 84 条真实样本",
    )
    s03 = 0.30 * step + 0.70 * nasa_rev
    add_para(
        doc,
        f"随 α 由 0.30 增至 0.70，84 条样本的 S 均值由 {s03.mean():.3f} 升至 {s07.mean():.3f}，"
        f"标准差由 {s03.std(ddof=1):.3f} 增至 {s07.std(ddof=1):.3f}，取值范围由 "
        f"{s03.min():.2f}–{s03.max():.2f} 展宽至 {s07.min():.2f}–{s07.max():.2f}。"
        "也就是说，把权重更多地压在客观步骤分上，整体绩效水平抬高，同时被试间的区分度也更大。",
    )

    add_heading(doc, "三、α = 0.70 下算出的 S")
    add_para(
        doc,
        "下图从四个角度呈现正式口径下的 S：(a) 真值分布；(b) 按任务类型的步骤分与合成 S；"
        "(c) 每条样本的 S 随 α 的变化轨迹；(d) 四模态 27 维 NASA 公式法的五折交叉验证预测效果。",
    )
    add_figure(doc, FIG_DIR / "fig2_s_results_docx.png")
    add_caption(doc, "图 2　α = 0.70 下合成绩效 S 的分布、分任务表现、α 敏感性与预测效果")

    add_para(doc, "表 1　真值 S 的描述统计（α = 0.70，n = 84）", indent=False,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_table(
        doc,
        ["样本数", "最小值", "最大值", "均值", "标准差"],
        [["84", f"{s07.min():.2f}", f"{s07.max():.2f}",
          f"{s07.mean():.3f}", f"{s07.std(ddof=1):.3f}"]],
        [3.0, 3.0, 3.0, 3.0, 3.0],
    )

    g = (
        df.groupby("task")
        .agg(n=("S07", "size"), step=("weighted_step_score", "mean"), S=("S07", "mean"))
        .reindex(TASK_ORDER)
    )
    add_para(doc, "表 2　按任务类型的步骤分与合成 S（α = 0.70）", indent=False,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_table(
        doc,
        ["任务", "条数", "步骤分均值", "合成 S 均值"],
        [[t, str(int(g.loc[t, "n"])), f"{g.loc[t, 'step']:.3f}", f"{g.loc[t, 'S']:.3f}"]
         for t in TASK_ORDER],
        [3.5, 3.5, 4.5, 4.5],
    )
    add_para(
        doc,
        "任务 1 至 5_6 的步骤分与合成 S 同步单调下降，与预设难度递增的设计一致；"
        "任务 5 是唯一一处合成 S（0.600）略高于步骤分（0.576）的情形，"
        "说明该任务客观完成度偏低但主观负荷并不重。",
    )

    add_heading(doc, "四、预测结果")
    add_para(
        doc,
        f"在四模态 27 维定额特征（眼动 6、脑电 5、心率 4、行为 12）、按被试 GroupKFold 五折的协议下，"
        f"先用浅树 XGBoost 预测 NASA-TLX，再按同一 α = 0.70 公式合成 S，"
        f"合成 S 的 pooled R² = {pred_r2:.3f}、MAE = {pred_mae:.3f}（图 2d），"
        f"中间量 NASA 的 R² = 0.553、MAE = 0.825。需要说明的是，S 的 70% 来自真实步骤记录、"
        f"不经过模型，因此评估生理与行为模型本身的能力应以 NASA 的 R² 为准。",
    )

    add_heading(doc, "五、可直接引用的一段话")
    add_para(
        doc,
        f"综合绩效 S 定义为客观步骤分与 NASA-TLX 反向分的加权合成，S = α × 步骤分 + (1 − α) × NASA 反向分，"
        f"正式口径取 α = 0.70。在 26 名被试、84 次任务上，真值 S 的范围为 "
        f"{s07.min():.2f}–{s07.max():.2f}，均值 {s07.mean():.3f}，标准差 {s07.std(ddof=1):.3f}。"
        f"按任务类型，步骤分均值由任务 1 的 0.896 单调降至任务 5_6 的 0.553，"
        f"合成 S 相应由 0.766 降至 0.492，与预设难度递增一致。",
    )

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print("[docx] 写出", build())
