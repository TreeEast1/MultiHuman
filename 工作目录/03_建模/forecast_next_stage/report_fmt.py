#!/usr/bin/env python3
"""Thesis/report Word styles matching the user's template.

一级标题 Heading 1：黑体 / TNR，小三 15pt，字符间距按二号，行距 2.41 倍，段前段后 12 磅
二级标题 Heading 2：黑体 / TNR，四号 14pt，行距 1.73 倍，段前段后 6 磅
三级标题 Heading 3：宋体 / TNR，小四 12pt，固定值 22 磅，段前段后 0.5 行
正文：宋体 / TNR，小四 12pt，两端对齐，固定值 22 磅，首行缩进 2 字符
图注：五号 10.5pt，居中，单倍行距
目录：宋体小四；一级不缩进，二级缩进 2 字，三级缩进 4 字
公式：Word OMML（LaTeX 公式编辑器可识别的 oMath），段落用图注样式
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Emu, Pt, RGBColor, Twips

HEI = "黑体"
SONG = "宋体"
TNR = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _rfonts(rPr, east, west=TNR):
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), west)
    rFonts.set(qn("w:hAnsi"), west)
    rFonts.set(qn("w:eastAsia"), east)
    rFonts.set(qn("w:cs"), west)
    return rFonts


def set_run_font(run, east, size, *, bold=False, west=TNR, spacing_pt=None):
    run.bold = bold
    run.italic = False
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.font.name = west
    rPr = run._element.get_or_add_rPr()
    _rfonts(rPr, east, west)
    if spacing_pt is not None:
        sp = rPr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            rPr.append(sp)
        sp.set(qn("w:val"), str(int(round(spacing_pt * 100))))  # hundredths of a point
    else:
        sp = rPr.find(qn("w:spacing"))
        if sp is not None:
            rPr.remove(sp)


def _style_rpr(style, east, size, *, bold=False, spacing_pt=None):
    style.font.name = TNR
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    rPr = style.element.get_or_add_rPr()
    _rfonts(rPr, east)
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), str(int(size * 2)))
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(int(size * 2)))
    if bold:
        b = rPr.find(qn("w:b"))
        if b is None:
            rPr.append(OxmlElement("w:b"))
        bCs = rPr.find(qn("w:bCs"))
        if bCs is None:
            rPr.append(OxmlElement("w:bCs"))
    else:
        for tag in ("w:b", "w:bCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn("w:val"), "0")
    if spacing_pt is not None:
        sp = rPr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            rPr.append(sp)
        sp.set(qn("w:val"), str(int(round(spacing_pt * 100))))


def _ppr(style, *, align=None, before=0, after=0, line=None, line_rule="auto",
         first_line=None, left=None, outline=None, keep=True, widow=True, single=False):
    pPr = style.element.get_or_add_pPr()

    def set_child(tag, **attrs):
        el = pPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            pPr.append(el)
        for k, v in attrs.items():
            el.set(qn(k), str(v))
        return el

    if align is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), align)

    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:before"), str(int(before * 20)))  # pt -> twips
    sp.set(qn("w:after"), str(int(after * 20)))
    if single:
        sp.set(qn("w:line"), "240")
        sp.set(qn("w:lineRule"), "auto")
    elif line is not None:
        if line_rule == "exact":
            sp.set(qn("w:line"), str(int(line * 20)))  # pt -> twips
            sp.set(qn("w:lineRule"), "exact")
        else:
            sp.set(qn("w:line"), str(int(round(line * 240))))
            sp.set(qn("w:lineRule"), "auto")

    ind = pPr.find(qn("w:ind"))
    if first_line is not None or left is not None:
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        if first_line is not None:
            ind.set(qn("w:firstLine"), str(int(first_line * 20)))
        if left is not None:
            ind.set(qn("w:left"), str(int(left * 20)))

    if outline is not None:
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is None:
            ol = OxmlElement("w:outlineLvl")
            pPr.append(ol)
        ol.set(qn("w:val"), str(outline))

    for tag, on in (("w:keepNext", keep), ("w:keepLines", keep)):
        el = pPr.find(qn(tag))
        if on:
            if el is None:
                pPr.append(OxmlElement(tag))
        elif el is not None:
            pPr.remove(el)

    if widow:
        wdw = pPr.find(qn("w:widowControl"))
        if wdw is None:
            pPr.append(OxmlElement("w:widowControl"))


def configure_styles(doc: Document):
    # 正文
    normal = doc.styles["Normal"]
    _style_rpr(normal, SONG, 12, bold=False)
    _ppr(normal, align="both", before=0, after=0, line=22, line_rule="exact",
         first_line=24, keep=False, widow=True)  # 2 chars at 12pt = 24pt

    h1 = doc.styles["Heading 1"]
    _style_rpr(h1, HEI, 15, bold=True)
    _ppr(h1, align="left", before=12, after=12, line=2.41, line_rule="auto",
         first_line=0, left=0, outline=0, keep=True)

    h2 = doc.styles["Heading 2"]
    _style_rpr(h2, HEI, 14, bold=True)
    _ppr(h2, align="left", before=6, after=6, line=1.73, line_rule="auto",
         first_line=0, left=0, outline=1, keep=True)

    h3 = doc.styles["Heading 3"]
    _style_rpr(h3, SONG, 12, bold=False)
    _ppr(h3, align="left", before=11, after=11, line=22, line_rule="exact",
         first_line=0, left=0, outline=2, keep=True)

    # 图：专用于嵌入型图片，单倍行距、无首行缩进，避免正文「固定 22 磅」把图裁成空框
    try:
        fig_st = doc.styles.add_style("图", 1)
    except ValueError:
        fig_st = doc.styles["图"]
    _style_rpr(fig_st, SONG, 12, bold=False)
    _ppr(fig_st, align="center", before=6, after=2, single=True,
         first_line=0, left=0, keep=True, widow=True)
    _clear_based_on(fig_st)
    _force_ind_chars(fig_st)

    # 图注
    try:
        cap = doc.styles.add_style("图注", 1)  # paragraph style
    except ValueError:
        cap = doc.styles["图注"]
    _style_rpr(cap, SONG, 10.5, bold=False)
    _ppr(cap, align="center", before=3, after=6, single=True,
         first_line=0, left=0, keep=False, widow=True)
    _clear_based_on(cap)
    _force_ind_chars(cap)

    for name, left_pt in (("TOC 1", 0), ("TOC 2", 24), ("TOC 3", 48)):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        _style_rpr(st, SONG, 12, bold=False)
        _ppr(st, align="left", before=0, after=0, line=22, line_rule="exact",
             first_line=0, left=left_pt, keep=False)

    # default page
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(1.5)


def add_run(p, text, east, size, *, bold=False, spacing_pt=None):
    run = p.add_run(text)
    set_run_font(run, east, size, bold=bold, spacing_pt=spacing_pt)
    return run


def add_toc_heading(doc):
    """「目录」二字居中，黑体小三，不纳入一级标题编号。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 2.41
    add_run(p, "目录", HEI, 15, bold=True)
    return p


def _clear_based_on(style):
    el = style.element.find(qn("w:basedOn"))
    if el is not None:
        style.element.remove(el)


def _force_ind_chars(style):
    pPr = style.element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:firstLineChars"), "0")
    ind.set(qn("w:left"), "0")


def add_page_break(doc):
    """分页符放在独立段，但不继承正文首行缩进，避免目录末出现空框。"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    for child in list(pPr):
        pPr.remove(child)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:firstLineChars"), "0")
    ind.set(qn("w:left"), "0")
    pPr.append(ind)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def add_h(doc, text, level):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    for run in p.runs:
        if level in (1, 2):
            set_run_font(run, HEI, {1: 15, 2: 14}[level], bold=True)
        else:
            set_run_font(run, SONG, 12, bold=False)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # style already created an empty run sometimes
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], SONG, 12)
    else:
        add_run(p, text, SONG, 12)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="图注")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # style may not apply run font until we set it
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], SONG, 10.5)
    else:
        add_run(p, text, SONG, 10.5)
    return p


def add_note(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, text, SONG, 10.5)
    return p


def add_toc_line(doc, text, level):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(24 * (level - 1))
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, text, SONG, 12)
    return p


def _force_embedded_inline(run):
    """把 python-docx 的不完整 drawing 改写成 Word 原生「嵌入型」wp:inline。

    缺 dist / effectExtent / picLocks / avLst 时，Word for Mac 常把图显示成
    带句柄的空框，文字环绕里看起来像「浮于文字上方」。
    """
    r = run._r
    drawing = r.find(qn("w:drawing"))
    if drawing is None:
        return
    inline = drawing.find(qn("wp:inline"))
    anchor = drawing.find(qn("wp:anchor"))
    src = inline if inline is not None else anchor
    if src is None:
        return
    extent = src.find(qn("wp:extent"))
    cx, cy = extent.get("cx"), extent.get("cy")
    docPr = src.find(qn("wp:docPr"))
    doc_id = docPr.get("id")
    name = docPr.get("name") or "Picture"
    blip = src.find(f".//{qn('a:blip')}")
    embed = blip.get(qn("r:embed"))
    nv = src.find(f".//{qn('pic:cNvPr')}")
    pic_name = nv.get("name") if nv is not None else name

    xml = (
        f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        f'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        f'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{doc_id}" name="{name}"/>'
        f'<wp:cNvGraphicFramePr>'
        f'<a:graphicFrameLocks noChangeAspect="1"/>'
        f'</wp:cNvGraphicFramePr>'
        f'<a:graphic>'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic>'
        f'<pic:nvPicPr>'
        f'<pic:cNvPr id="0" name="{pic_name}"/>'
        f'<pic:cNvPicPr><a:picLocks noChangeAspect="1" noChangeArrowheads="1"/></pic:cNvPicPr>'
        f'</pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip r:embed="{embed}" cstate="print"/>'
        f'<a:stretch><a:fillRect/></a:stretch>'
        f'</pic:blipFill>'
        f'<pic:spPr bwMode="auto">'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'<a:noFill/><a:ln><a:noFill/></a:ln>'
        f'</pic:spPr>'
        f'</pic:pic>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:inline>'
        f'</w:drawing>'
    )
    r.remove(drawing)
    r.append(parse_xml(xml))
    rPr = r.get_or_add_rPr()
    if rPr.find(qn("w:noProof")) is None:
        rPr.append(OxmlElement("w:noProof"))
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), "24")


def add_figure(doc, path, caption, width_cm=14.5):
    p = doc.add_paragraph(style="图")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:before"), "120")
    sp.set(qn("w:after"), "40")
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "auto")
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:firstLineChars"), "0")
    ind.set(qn("w:left"), "0")
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "center")
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    _force_embedded_inline(run)
    add_caption(doc, caption)


def _omath_t(text):
    r = OxmlElement("m:r")
    rPr = OxmlElement("m:rPr")
    sty = OxmlElement("m:sty")
    sty.set(qn("m:val"), "p")
    rPr.append(sty)
    r.append(rPr)
    t = OxmlElement("m:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _omath_sub(base, sub):
    sSub = OxmlElement("m:sSub")
    e = OxmlElement("m:e")
    e.append(_omath_t(base))
    sSub.append(e)
    sube = OxmlElement("m:sub")
    sube.append(_omath_t(sub))
    sSub.append(sube)
    return sSub


def _omath_frac(num, den):
    f = OxmlElement("m:f")
    num_e = OxmlElement("m:num")
    num_e.append(_omath_t(num))
    den_e = OxmlElement("m:den")
    den_e.append(_omath_t(den))
    f.append(num_e)
    f.append(den_e)
    return f


def add_equation(doc, children):
    """Insert a centered OMML paragraph using 图注 style (Word 公式编辑器可再编辑)."""
    p = doc.add_paragraph(style="图注")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    oMathPara = OxmlElement("m:oMathPara")
    pr = OxmlElement("m:oMathParaPr")
    jc = OxmlElement("m:jc")
    jc.set(qn("m:val"), "center")
    pr.append(jc)
    oMathPara.append(pr)
    oMath = OxmlElement("m:oMath")
    for ch in children:
        oMath.append(ch)
    oMathPara.append(oMath)
    p._p.append(oMathPara)
    return p


def eq_s_step():
    return [
        _omath_sub("S", "step"),
        _omath_t(" = 0.75 "),
        _omath_sub("r", "key"),
        _omath_t(" + 0.25 "),
        _omath_sub("r", "nkey"),
    ]


def eq_l_rev():
    return [
        _omath_sub("L", "rev"),
        _omath_t(" = 1 - "),
        _omath_frac("L", "10"),
    ]


def eq_s():
    return [
        _omath_t("S = 0.70 "),
        _omath_sub("S", "step"),
        _omath_t(" + 0.30 "),
        _omath_sub("L", "rev"),
    ]


def eq_s_hat():
    return [
        _omath_sub("S", "hat"),
        _omath_t(" = 0.70 "),
        _omath_sub("S", "step"),
        _omath_t(" + 0.30 (1 - "),
        _omath_frac(r"L̂", "10"),
        _omath_t(")"),
    ]


def eq_alpha():
    return [
        _omath_t("S("),
        _omath_t("α"),
        _omath_t(") = α "),
        _omath_sub("S", "step"),
        _omath_t(" + (1 - α) "),
        _omath_sub("L", "rev"),
    ]


def set_cell_borders(cell, sz="8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(sum(widths_cm) * 567)))
    tblW.set(qn("w:type"), "dxa")
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


def fill_cell(cell, text, *, header=False, align="center", size=10.5):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    east = HEI if header else SONG
    add_run(p, str(text), east, size, bold=header)
    set_cell_borders(cell)
    set_cell_margins(cell)


def make_table(doc, headers, rows, widths, *, left_cols=None, size=10.5):
    left_cols = left_cols or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed(table, widths)
    for j, h in enumerate(headers):
        fill_cell(table.rows[0].cells[j], h, header=True, size=size)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            fill_cell(
                table.rows[i + 1].cells[j],
                val,
                header=False,
                align="left" if j in left_cols else "center",
                size=size,
            )
    for i, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))
        if i == 0 and trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))
    return table


def add_header_footer(doc, header_text):
    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.first_line_indent = Pt(0)
        add_run(hp, header_text, SONG, 9)
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.first_line_indent = Pt(0)
        add_run(fp, "— ", SONG, 10.5)
        run = fp.add_run()
        set_run_font(run, SONG, 10.5)
        fld1 = OxmlElement("w:fldChar")
        fld1.set(qn("w:fldCharType"), "begin")
        run._r.append(fld1)
        run2 = fp.add_run()
        set_run_font(run2, SONG, 10.5)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run2._r.append(instr)
        run3 = fp.add_run()
        set_run_font(run3, SONG, 10.5)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        run3._r.append(fld2)
        add_run(fp, " —", SONG, 10.5)


def new_doc(header_text) -> Document:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, header_text)
    # ensure OMML namespace
    root = doc.element
    if "xmlns:m" not in root.xml[:500]:
        root.set(qn("xmlns:m"), MATH_NS)
    return doc


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    add_run(p, title, HEI, 18, bold=True)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.first_line_indent = Pt(0)
        p2.paragraph_format.space_after = Pt(12)
        add_run(p2, subtitle, SONG, 14)
