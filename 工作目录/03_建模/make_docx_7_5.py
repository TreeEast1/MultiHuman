#!/usr/bin/env python3
"""按原报告体例重写 7.5 结果分析小结。

口径与《CJL改-V2徐批注-研究内容报告_V4》第7.3、7.4节及正式实验一致：
26 被试、84 条任务；S＝0.70×步骤分＋0.30×（1−NASA/10）。
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SRC = Path("/Users/licochen/Desktop/CJL改-V2徐批注-研究内容报告_V4.docx")
HERE = Path(__file__).resolve().parent
OUT_PROJ = HERE / "7.5结果分析小结.docx"
OUT_DESK = Path("/Users/licochen/Desktop/7.5结果分析小结.docx")

HEI = "黑体"
SONG = "宋体"
TNR = "Times New Roman"
BLACK = RGBColor(0, 0, 0)


def _rfonts(rPr, east: str, west: str = TNR) -> None:
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), west)
    rFonts.set(qn("w:hAnsi"), west)
    rFonts.set(qn("w:eastAsia"), east)
    rFonts.set(qn("w:cs"), west)
    rFonts.set(qn("w:hint"), "eastAsia")


def set_run_font(run, east: str, size: float, *, bold: bool = False) -> None:
    run.bold = bold
    run.italic = False
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.font.name = TNR
    rPr = run._element.get_or_add_rPr()
    _rfonts(rPr, east)
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), str(int(size * 2)))
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(int(size * 2)))


def _ppr_exact(p, *, first_line: bool, center: bool = False,
               before: int = 0, after: int = 0) -> None:
    """与原文 7.4/7.5 直接格式一致：固定行距 22 磅（440 twips）。"""
    pf = p.paragraph_format
    pf.line_spacing = Pt(22)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(before / 20)
    pf.space_after = Pt(after / 20)
    pf.widow_control = False
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if first_line:
        ind.set(qn("w:firstLine"), "480")
        ind.set(qn("w:firstLineChars"), "200")
    else:
        ind.set(qn("w:firstLine"), "0")
        ind.set(qn("w:firstLineChars"), "0")
    if center:
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "center")


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.clear()
    _ppr_exact(p, first_line=False, before=120, after=120)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._element.get_or_add_pPr()
    jc = pPr.find(qn("w:jc"))
    if jc is not None:
        pPr.remove(jc)
    run = p.add_run(text)
    set_run_font(run, HEI, 14, bold=True)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.clear()
    _ppr_exact(p, first_line=True)
    run = p.add_run(text)
    set_run_font(run, SONG, 12)


def apply_src_page(doc: Document) -> None:
    """页边距、纸张与原报告末节一致。"""
    src = Document(str(SRC))
    src_sec = src.sections[-1]
    sec = doc.sections[0]
    sec.page_width = src_sec.page_width
    sec.page_height = src_sec.page_height
    sec.top_margin = src_sec.top_margin
    sec.bottom_margin = src_sec.bottom_margin
    sec.left_margin = src_sec.left_margin
    sec.right_margin = src_sec.right_margin
    sec.header_distance = src_sec.header_distance
    sec.footer_distance = src_sec.footer_distance
    sec.gutter = src_sec.gutter


def build() -> Path:
    doc = Document()
    apply_src_page(doc)
    normal = doc.styles["Normal"]
    normal.font.name = TNR
    normal.font.size = Pt(12)
    rPr = normal.element.get_or_add_rPr()
    _rfonts(rPr, SONG)
    pPr = normal.element.get_or_add_pPr()
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "both")

    add_h2(doc, "7.5 结果分析小结")

    add_body(
        doc,
        "（1）综合绩效评分S在本研究中被操作化为任务执行质量与认知负荷反向分量的加权合成指数，"
        "而非独立于实验设计的官方量表。步骤分由关键子任务完成率0.75与非关键完成率0.25加权得到，"
        "负荷反向分定义为1−NASA-TLX加权总分/10，正式配比为S＝0.70×步骤分＋0.30×负荷反向分。"
        "S升高对应于关键操作完成更为充分，或在相当执行水平下主观负荷更低；"
        "S降低则意味着步骤缺失、操作失误或认知资源耗竭中至少一类因素占优。"
        "据此，S既可用于跨被试、跨任务的绩效比较，也可反向分解绩效差异的来源，"
        "从而识别关键绩效因素与薄弱环节：前者刻画拉开综合绩效的主导信息，"
        "后者将差异定位至具体任务类型、失误步骤及负荷状态。",
    )

    add_body(
        doc,
        "（2）在26名被试、84次任务上，综合绩效的高低首先表现为任务难度与关键步骤完成率的分层，"
        "其次表现为同等步骤表现下认知负荷的分化。任务1与任务2的步骤分均值分别为0.896和0.850，"
        "合成S为0.766和0.742，构成高绩效任务类型；任务5_6步骤分仅0.553，"
        "NASA加权总分均值达6.492，合成S降至0.492，构成全样本最薄弱环节。"
        "按合成S三分位均分后，高绩效档（28条）步骤分均值0.966、关键子任务完成率1.000、"
        "NASA均值4.317、S均值0.847，表明高绩效以关键步骤近乎全部完成为前提，并伴随相对较低的主观负荷；"
        "低绩效档步骤分仅0.324、关键完成率0.288、NASA均值5.662、S均值0.357，"
        "步骤分相差0.642构成拉开S的主效应，负荷同时高出约1.35分，说明低绩效并非单纯的未完成，"
        "而常与更高的认知代价并存。日志层面的总体统计进一步给出失误类型与失误步骤："
        "被试01–12共36场标注中，平均完成率为0.889，35场出现错误动作，32场出现额外操作，27场存在缺失步骤；"
        "主要失误类型为出现额外操作（成功准则为顺序正确性）与参数不符合要求（集中于实验复位），"
        "高频缺失步骤为任务五的“微调RCV046VP”“微调RCV061VP”（各11场）"
        "以及任务三的“恢复RCP016KG、RCP001VP/002VP为初始状态”（7场）。"
        "以操纵员02为例，节点级分析中成功19项（65.52%）、失败10项（34.48%），"
        "失败集中于RCP001VP实验的额外操作与复位参数不合格，属于步骤失败型低绩效；"
        "其任务5_6步骤分已达0.800，但NASA为6.600，合成S仅为0.662，"
        "84条中同类“步骤尚可而负荷偏高”的样本共8例，属于高负荷隐匿型低绩效。"
        "由此，关键绩效因素可归结为关键子任务完成率以及支撑较低负荷的视觉搜索与操作过程；"
        "薄弱环节则集中于高难度任务5_6、复位类步骤、顺序性子任务中的额外操作，"
        "以及步骤完成尚可但负荷已显著升高的隐匿状态。",
    )

    add_body(
        doc,
        "（3）NASA-TLX负荷三分类结果表明，训练折内的模态互信息筛选将264维特征压缩至27维后，"
        "浅层XGBoost的Accuracy和Macro-F1均达到0.774，优于67维候选特征集的0.750和0.747。"
        "该结果表明，模态约束的紧凑特征选择不仅能够降低小样本条件下的冗余与过拟合风险，"
        "还能提高认知负荷模型的跨被试泛化能力。"
        "负荷识别同时表现出显著的多模态互补性与等级边界效应：眼动为最强单模态（Macro-F1＝0.699），"
        "眼动与行为融合后提升至0.740，完整四模态融合进一步达到0.774；"
        "模型对高负荷识别最佳（F1＝0.852），中负荷识别相对较弱（F1＝0.691），"
        "误判主要集中于两个分档切点附近。"
        "这说明视觉搜索与操作行为构成负荷识别的核心信息，高负荷状态可被稳定区分，"
        "从而为解释“步骤完成尚可但综合绩效不高”提供独立于日志计数的依据；"
        "中等负荷更接近具有较高不确定性的连续过渡状态，不宜单独作为薄弱环节的判据。",
    )

    add_body(
        doc,
        "（4）在已知任务步骤表现并利用多模态模型估计负荷分量的条件下，"
        "浅层XGBoost对综合绩效评分S的折外估计达到pooled R²＝0.979、MAE＝0.025；"
        "相较于均值负荷基线，MAE由0.040降至0.025，误差下降37.5%。"
        "五折R²均不低于0.933，表明紧凑多模态模型能够在不同被试之间稳定地修正仅依赖任务步骤表现的绩效估计。"
        "模态消融揭示了以眼动和行为为核心、心率和脑电提供边际补充的层级贡献结构："
        "眼动单模态已达R²＝0.968，眼动与行为融合提高至0.978，"
        "眼动、心率和行为三模态与完整四模态均达到0.979；脑电、心率单模态分别为0.949和0.944。"
        "据此，视觉注意与实际操作过程是跨被试绩效差异的主要信息来源，"
        "脑电和心率更适合作为辅助信息，而非独立的绩效评价依据。",
    )

    add_body(
        doc,
        "（5）失误类型与失误步骤可由传统操作日志统计得到，多模态方法的增量并不在于重复给出上述清单，"
        "而在于补足传统计数无法观测的负荷结构及其对综合绩效的修正。"
        "其一，在按被试划分的折外条件下，四模态27维输入使负荷三分类Macro-F1达到0.774、"
        "连续负荷估计R²达到0.553，从而使S中30%的负荷分量可在未见过的被试上估计，而无须依赖当场问卷。"
        "其二，眼动尤其是兴趣区注视结构构成最强单模态，将关键绩效因素由操作对错推进至视觉搜索是否覆盖正确区域、"
        "注意资源是否过度分散。"
        "其三，在步骤分已知的前提下，多模态估计使仅依赖步骤表现的绩效误差下降37.5%，"
        "得以识别步骤完成尚可但主观负荷已高的隐匿型低绩效。"
        "因此，传统方法回答的是操作序列中的错误步骤与失误类型，"
        "多模态方法进一步区分同等步骤完成率下的负荷代价，"
        "从而把绩效差异解释为“未完成”与“高负荷下完成”两类机制，并对准相应的薄弱环节。",
    )

    doc.save(OUT_PROJ)
    shutil.copy2(OUT_PROJ, OUT_DESK)
    print("wrote", OUT_PROJ)
    print("copy ", OUT_DESK)
    return OUT_PROJ


if __name__ == "__main__":
    build()
